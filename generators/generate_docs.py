#!/usr/bin/env python3
"""
CA書類自動生成スクリプト
Usage: python3 generate_docs.py <input_json_path>
Prints: JSON array of { label, filename, path }
"""

import sys, json, os, copy, shutil, warnings
from datetime import datetime

warnings.filterwarnings('ignore')

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), '..', 'templates')


# ─────────────────────────────────────────────
# ユーティリティ
# ─────────────────────────────────────────────

def stamp():
    return datetime.now().strftime('%Y%m%d_%H%M')

def safe_name(name):
    for ch in r'\/:*?"<>|': name = name.replace(ch, '_')
    return name


def format_interview_date_jp(raw_date):
    """'2026-04-28' → '4月28日' に整形。不正な値はそのまま返す。"""
    if not raw_date:
        return ''
    s = str(raw_date).strip()
    # YYYY-MM-DD or YYYY/MM/DD
    import re
    m = re.match(r'^\s*(\d{4})[-/](\d{1,2})[-/](\d{1,2})', s)
    if m:
        return f"{int(m.group(2))}月{int(m.group(3))}日"
    # MM/DD
    m = re.match(r'^\s*(\d{1,2})[-/](\d{1,2})', s)
    if m:
        return f"{int(m.group(1))}月{int(m.group(2))}日"
    return s

def format_interview_time(iv):
    """面接時間帯のセル表示を決定。
    - note に「終日」等があればその文字列を優先
    - start/end が両方空 → '終日可'
    - 09:00〜18:00 → '終日可'
    - 片方のみ → 'HH:MM〜' or '〜HH:MM'
    - それ以外 → 'HH:MM〜HH:MM'
    """
    ts = (iv.get('timeStart') or '').strip()
    te = (iv.get('timeEnd') or '').strip()
    slot = (iv.get('timeSlot') or '').strip()
    note = (iv.get('note') or '').strip()
    # 明示的 note 優先
    if note and ('終日' in note or '午前' in note or '午後' in note):
        return note
    # 両方空 → 終日
    if not ts and not te:
        return '終日可' if note == '' else note
    # 09:00〜18:00 は終日とみなす
    if ts in ('09:00','9:00') and te in ('18:00','17:00','17:30','19:00','19:45'):
        return '終日可'
    if ts and te:
        return f"{ts}〜{te}"
    if ts and not te:
        return f"{ts}〜"
    if not ts and te:
        return f"〜{te}"
    return slot


def to_hiragana(s):
    """カタカナ→ひらがな変換。既にひらがな/ascii はそのまま。"""
    if not s:
        return s
    out = []
    for ch in s:
        code = ord(ch)
        # カタカナ（U+30A1〜U+30F6）を ひらがな（U+3041〜U+3096）へ
        if 0x30A1 <= code <= 0x30F6:
            out.append(chr(code - 0x60))
        else:
            out.append(ch)
    return ''.join(out)

def _get_template_rpr(para):
    """段落内の既存Runから rPr（書式）を探して deepcopy で返す。無ければ None。"""
    from docx.oxml.ns import qn
    for run in para.runs:
        rpr = run._element.find(qn('w:rPr'))
        if rpr is not None:
            return copy.deepcopy(rpr)
    # 段落スタイル側の rPr を探索（pPr/rPr）
    ppr = para._p.find(qn('w:pPr'))
    if ppr is not None:
        rpr = ppr.find(qn('w:rPr'))
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None

def _apply_rpr(run, rpr):
    """run に rPr を適用（既存の rPr は置換）"""
    from docx.oxml.ns import qn
    if rpr is None:
        return
    old = run._element.find(qn('w:rPr'))
    if old is not None:
        run._element.remove(old)
    run._element.insert(0, copy.deepcopy(rpr))

def set_paragraph_text(para, text):
    """段落の文字列を置換（書式を先頭Runから継承。空段落でもテンプレ書式を維持）"""
    from docx.oxml.ns import qn
    tmpl_rpr = _get_template_rpr(para)
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = text
        if para.runs[0]._element.find(qn('w:rPr')) is None and tmpl_rpr is not None:
            _apply_rpr(para.runs[0], tmpl_rpr)
    else:
        new_run = para.add_run(text)
        if tmpl_rpr is not None:
            _apply_rpr(new_run, tmpl_rpr)

def add_paragraph_after(doc_body, ref_element, text='', style=None, ref_para=None):
    """ref_element の直後に段落を挿入。ref_para の書式を継承。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    new_p = OxmlElement('w:p')
    # 参照段落の pPr（段落書式）をコピー
    if ref_para is not None:
        src_ppr = ref_para._p.find(qn('w:pPr'))
        if src_ppr is not None:
            new_p.append(copy.deepcopy(src_ppr))
    r = OxmlElement('w:r')
    # 参照段落から rPr（文字書式）をコピー
    if ref_para is not None:
        tmpl_rpr = _get_template_rpr(ref_para)
        if tmpl_rpr is not None:
            r.append(copy.deepcopy(tmpl_rpr))
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    new_p.append(r)
    ref_element.addnext(new_p)
    return new_p

def _clear_paragraph(p):
    """段落内の全runのテキストをクリア（書式は維持）"""
    for run in p.runs:
        run.text = ''

def _write_paragraph(p, text, tmpl_rpr=None):
    """段落に1行テキストを書込む。既存runがあれば書式維持。"""
    from docx.oxml.ns import qn
    val = str(text) if text is not None else ''
    _clear_paragraph(p)
    if p.runs:
        p.runs[0].text = val
        if p.runs[0]._element.find(qn('w:rPr')) is None and tmpl_rpr is not None:
            _apply_rpr(p.runs[0], tmpl_rpr)
    else:
        new_run = p.add_run(val)
        if tmpl_rpr is not None:
            _apply_rpr(new_run, tmpl_rpr)

def cell_write(cell, text):
    """セルの全段落をクリアして書き直す。
    textに \n が含まれる場合は各行を別段落として書き、
    テンプレの段落数より多ければ追加、少なければ残りは空にする。
    フォント書式は元の段落から継承する。"""
    from docx.oxml.ns import qn
    val = str(text) if text is not None else ''
    lines = val.split('\n')
    p_list = list(cell.paragraphs)
    if not p_list:
        cell.add_paragraph('')
        p_list = list(cell.paragraphs)

    # 先頭段落の書式を保持（以降の段落複製用）
    tmpl_rpr = _get_template_rpr(p_list[0])

    # 行数と段落数を合わせる
    while len(p_list) < len(lines):
        # 先頭段落を複製してセルに追加
        new_p = copy.deepcopy(p_list[0]._p)
        # 中身クリア
        for r in new_p.findall(qn('w:r')):
            for t in r.findall(qn('w:t')):
                t.text = ''
        p_list[-1]._p.addnext(new_p)
        # 再取得
        p_list = list(cell.paragraphs)

    # 各段落に書き込み（余った段落は空文字）
    for i, p in enumerate(p_list):
        line = lines[i] if i < len(lines) else ''
        _write_paragraph(p, line, tmpl_rpr)


def xlsx_write(ws, cell_ref, value):
    """xlsx セルへ書き込み（結合セルの場合は左上セルが target）"""
    ws[cell_ref] = value


def tc(table, r, c):
    """python-docx の Table.cell(r,c) は列幅不揃いのテーブルでズレるため、
    rows[r].cells[c] 経由で安全にアクセスする。"""
    return table.rows[r].cells[c]


def set_cell_vertical_top(cell):
    """セルの垂直位置を「上揃え」に設定（縦マージセルでの中央寄せ回避）"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    # 既存 vAlign を削除
    existing = tcPr.find(qn('w:vAlign'))
    if existing is not None:
        tcPr.remove(existing)
    v = OxmlElement('w:vAlign')
    v.set(qn('w:val'), 'top')
    tcPr.append(v)


def force_font(doc, font_name='ＭＳ 明朝'):
    """文書内の全Runに対してフォントを強制適用（ASCII/hAnsi/eastAsia/cs すべて）。
    Wordの既定フォントも上書きして、後から追加される文字も同じフォントになるようにする。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def apply_rfonts(rPr):
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        for attr in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
            rFonts.set(qn(f'w:{attr}'), font_name)

    def walk_element_runs(el):
        """任意の要素配下の全 w:r を列挙"""
        return el.iter(qn('w:r'))

    # 1) 本文段落・ヘッダ・フッタ・テーブル…全ての run に適用
    body = doc.element.body
    for r in walk_element_runs(body):
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
        apply_rfonts(rPr)

    # 2) スタイル定義(styles.xml)の Normal / Default などにも適用 → 以降追加される文字も明朝に
    try:
        styles_el = doc.styles.element
        for rPr in styles_el.iter(qn('w:rPr')):
            apply_rfonts(rPr)
        # docDefaults のデフォルト run プロパティ
        for rPrDefault in styles_el.iter(qn('w:rPrDefault')):
            rPr = rPrDefault.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                rPrDefault.insert(0, rPr)
            apply_rfonts(rPr)
    except Exception:
        pass


def _compact_gender_cell(cell):
    """性別値セル：先頭段落の後にある空段落を削除して縦余白を減らす"""
    from docx.oxml.ns import qn
    paras = list(cell.paragraphs)
    # 先頭は保持、残りの空段落は削除（内容ありなら保持）
    for p in paras[1:]:
        if not p.text.strip():
            p._p.getparent().remove(p._p)


def set_paragraph_alignment(para, alignment):
    """段落の水平揃えを設定。alignment: 'left','center','right','both'"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._p.insert(0, pPr)
    existing = pPr.find(qn('w:jc'))
    if existing is not None:
        pPr.remove(existing)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), alignment)
    pPr.append(jc)


# ─────────────────────────────────────────────
# 1. 汎用 履歴書.docx
# ─────────────────────────────────────────────

def gen_resume(d, out_dir, sid):
    from docx import Document
    import copy

    src = os.path.join(TEMPLATES_DIR, '履歴書.docx')
    dst = os.path.join(out_dir, f"{safe_name(d.get('candidateName','候補者'))}_履歴書_{stamp()}.docx")
    shutil.copy2(src, dst)
    doc = Document(dst)

    # ── 作成日：テンプレ既存の「YYYY年M月D日 現在」を今日の日付に差し替え ──
    today = datetime.now()
    today_str = f"{today.year}年{today.month}月{today.day}日 現在"
    import re
    date_pat = re.compile(r'\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日\s*現在')
    for p in doc.paragraphs:
        if date_pat.search(p.text):
            # 段落内の全runテキストから旧日付を置換
            # まずテキスト全体を取得→置換→先頭runに書き戻し、残りrunをクリア
            full_text = p.text
            new_text = date_pat.sub(today_str, full_text)
            set_paragraph_text(p, new_text)

    # ── Table 0: 個人情報 ──
    t0 = doc.tables[0]
    # row0: [0,1]=ふりがな（ひらがな化）
    cell_write(tc(t0, 0, 1), to_hiragana(d.get('candidateNameKana', '')))
    # 性別値は「性別ラベル [0,2]」の真下 [1,2] に書く（[0,3] は写真領域なので使わない）
    gender_cell = tc(t0, 1, 2)
    gender_raw = (d.get('gender', '') or '').strip()
    # 「女性」→「女」、「男性」→「男」、それ以外は先頭1文字
    gender_short = '女' if '女' in gender_raw else ('男' if '男' in gender_raw else (gender_raw[:1] if gender_raw else ''))
    cell_write(gender_cell, gender_short)
    # セル内 中央揃え（水平・垂直とも中央）
    from docx.enum.table import WD_ALIGN_VERTICAL
    gender_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for p in gender_cell.paragraphs:
        set_paragraph_alignment(p, 'center')
    # row1: [1,1]=氏名
    cell_write(tc(t0, 1, 1), d.get('candidateName', ''))
    # row2: [2,1]=生年月日
    by = d.get('birthYear', ''); bm = d.get('birthMonth', ''); bd = d.get('birthDay', '')
    age = d.get('age', '')
    bdate_str = f"{by}年{bm}月{bd}日生　（ 満{age}歳 ）" if by else ''
    cell_write(tc(t0, 2, 1), bdate_str)
    # row3: [3,1]=住所ふりがな（ひらがな化）, [3,3]=自宅電話
    cell_write(tc(t0, 3, 1), to_hiragana(d.get('addressKana', '')))
    cell_write(tc(t0, 3, 3), d.get('phone', ''))
    # row4: [4,1]=〒+住所（番地含むフル住所）
    zip_str = (d.get('zipCode','') or '').strip()
    addr_full = (d.get('address','') or '').strip()
    zip_disp = f"〒{zip_str}" if zip_str else ''
    cell_write(tc(t0, 4, 1), f"{zip_disp}\n{addr_full}" if zip_disp else addr_full)
    # row5: [5,3]=携帯電話
    cell_write(tc(t0, 5, 3), d.get('mobile', d.get('phone', '')))
    # row6: [6,1]=email
    cell_write(tc(t0, 6, 1), d.get('email', ''))

    # ── Table 1: 学歴・職歴（row0=ヘッダー, row1=学歴ラベル, rows2-5=学歴, row6=職歴ラベル, rows7-20=職歴, row22=免許資格ヘッダー, rows23-26=資格）
    t1 = doc.tables[1]

    # 学歴（古い順→新しい順）
    edu = list(d.get('educationHistory', []) or [])
    def edu_key(e):
        try: return (int(e.get('year') or 0), int(e.get('month') or 0))
        except: return (0, 0)
    edu.sort(key=edu_key)
    for i, e in enumerate(edu[:4]):
        row = 2 + i
        cell_write(tc(t1, row, 0), str(e.get('year', '') or ''))
        cell_write(tc(t1, row, 1), str(e.get('month', '') or ''))
        cell_write(tc(t1, row, 2), e.get('content', ''))

    # 職歴
    works = d.get('workHistory', [])
    wrow = 7
    for w in works:
        if wrow > 20: break
        company = w.get('company', '')
        etype   = w.get('employmentType', '正社員')
        sy, sm  = w.get('startYear', ''), w.get('startMonth', '')
        ey, em  = w.get('endYear', ''),   w.get('endMonth', '')
        current = w.get('isCurrent', False)

        # 入社行
        cell_write(tc(t1, wrow, 0), str(sy))
        cell_write(tc(t1, wrow, 1), str(sm))
        cell_write(tc(t1, wrow, 2), f"{company}　入社（{etype}）")
        wrow += 1
        if wrow > 20: break

        # 退社 or 現在
        cell_write(tc(t1, wrow, 0), str(ey) if not current else '')
        cell_write(tc(t1, wrow, 1), str(em) if not current else '')
        cell_write(tc(t1, wrow, 2), f"{company}　退社" if not current else '現在に至る')
        wrow += 1
        if wrow > 20: break

    # 最終職歴が退社済み（isCurrent=False）の場合も、末尾に「現在に至る」を追加
    if works and not works[-1].get('isCurrent', False) and wrow <= 20:
        cell_write(tc(t1, wrow, 2), '現在に至る')
        wrow += 1

    # 以上（右寄せ）
    if wrow <= 20:
        izyou_cell = tc(t1, wrow, 2)
        cell_write(izyou_cell, '以上')
        for p in izyou_cell.paragraphs:
            if p.text.strip():
                set_paragraph_alignment(p, 'right')

    # 免許・資格
    # licenseHistory（年月付き）が優先、なければ licenses（名前のみ）
    license_hist = list(d.get('licenseHistory', []) or [])
    licenses_plain = list(d.get('licenses', []) or [])

    if license_hist:
        # 時系列昇順
        def lic_key(x):
            try: return (int(x.get('year') or 0), int(x.get('month') or 0))
            except: return (0, 0)
        license_hist.sort(key=lic_key)
        # 運転免許がなければデフォルト追加
        has_driver = any(('運転' in (x.get('name','') or '') or '普通自動車' in (x.get('name','') or '')) for x in license_hist)
        if not has_driver:
            license_hist.insert(0, {'year': '', 'month': '', 'name': '普通自動車第一種運転免許（AT限定/MT）'})
        for i, lic in enumerate(license_hist[:4]):
            row = 23 + i
            try:
                cell_write(tc(t1, row, 0), str(lic.get('year','') or ''))
                cell_write(tc(t1, row, 1), str(lic.get('month','') or ''))
                cell_write(tc(t1, row, 2), lic.get('name',''))
            except Exception:
                pass
    else:
        has_driver = any(('運転' in (l or '') or '普通自動車' in (l or '')) for l in licenses_plain)
        if not has_driver:
            licenses_plain.insert(0, '普通自動車第一種運転免許（AT限定/MT）')
        for i, lic in enumerate(licenses_plain[:4]):
            row = 23 + i
            try:
                cell_write(tc(t1, row, 2), lic)
            except Exception:
                pass

    # ── Table 2: 通勤情報 ──
    t2 = doc.tables[2]
    line_raw = (d.get('nearestLine', '') or '').strip()
    sta_raw  = (d.get('nearestStation', '') or '').strip()
    # 末尾の「線」「駅」を重複させない
    line = line_raw[:-1] if line_raw.endswith('線') else line_raw
    sta  = sta_raw[:-1]  if sta_raw.endswith('駅')  else sta_raw
    ctime = str(d.get('commuteTime', '') or '').strip()
    # 末尾の「分」「時間」「分間」などを除去して数値部分だけ残す
    ctime_clean = ctime
    for suf in ('時間', '分間', '分'):
        if ctime_clean.endswith(suf):
            ctime_clean = ctime_clean[:-len(suf)].strip()
            break
    ctime_str = f"　　約　　{ctime_clean}　分" if ctime_clean else "　　約　　　　分"
    cell_write(tc(t2, 0, 0), f"通勤時間\n{ctime_str}")
    if line or sta:
        cell_write(tc(t2, 1, 0), f"最寄駅\u3000\n{line}線\u3000{sta}駅")
    else:
        cell_write(tc(t2, 1, 0), "最寄駅\u3000")

    # ── Table 3: 本人希望記入欄 ──
    # ラベル行は残し、それ以外の段落（既存「貴社規定に従います。」等）を削除／空化
    # jobPreference があれば本文行に書き込む
    t3 = doc.tables[3]
    t3_cell = tc(t3, 0, 0)
    hope_text = (d.get('jobPreference') or '').strip()
    written = False
    for i, p in enumerate(t3_cell.paragraphs):
        txt = p.text.strip()
        if '本人希望記入欄' in txt:
            continue
        if hope_text and not written:
            _write_paragraph(p, hope_text, _get_template_rpr(p))
            written = True
        else:
            _clear_paragraph(p)

    doc.save(dst)
    return {'label': '汎用 履歴書', 'filename': os.path.basename(dst), 'path': dst}


# ─────────────────────────────────────────────
# 2. 汎用 職務経歴書.docx
# ─────────────────────────────────────────────

def gen_career(d, out_dir, sid):
    from docx import Document
    from docx.oxml import OxmlElement
    from copy import deepcopy
    import lxml.etree as etree

    src = os.path.join(TEMPLATES_DIR, '職務経歴書.docx')
    dst = os.path.join(out_dir, f"{safe_name(d.get('candidateName','候補者'))}_職務経歴書_{stamp()}.docx")
    shutil.copy2(src, dst)
    doc = Document(dst)

    today = datetime.now()
    name  = d.get('candidateName', '')
    works = d.get('workHistory', [])

    # 本文ブロックを body 要素リストとして操作
    body = doc.element.body

    def get_body_elements():
        return list(body)

    def paras_and_tables():
        """body 内の段落・テーブルを (index, elem, type) で返す"""
        result = []
        for i, el in enumerate(body):
            t = el.tag.split('}')[-1]
            if t in ('p', 'tbl'):
                result.append((i, el, t))
        return result

    elems = get_body_elements()

    def find_para_with_text(needle):
        for el in elems:
            if el.tag.endswith('}p'):
                txt = ''.join(t.text or '' for t in el.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))
                if needle in txt:
                    return el
        return None

    def set_para_elem_text(el, text):
        ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        runs = el.findall(f'.//{{{ns}}}r')
        for r in runs:
            for t in r.findall(f'{{{ns}}}t'):
                t.text = ''
        if runs:
            t_el = runs[0].find(f'{{{ns}}}t')
            if t_el is None:
                t_el = OxmlElement('w:t')
                runs[0].append(t_el)
            t_el.text = text
        else:
            r_el = OxmlElement('w:r')
            t_el = OxmlElement('w:t')
            t_el.text = text
            r_el.append(t_el)
            el.append(r_el)

    def insert_paragraph_after(ref_el, text, copy_style_from=None):
        """ref_el の後ろに段落を追加してその elem を返す"""
        src_el = copy_style_from if copy_style_from is not None else ref_el
        new_p = deepcopy(src_el)
        set_para_elem_text(new_p, text)
        ref_el.addnext(new_p)
        return new_p

    # P[2]: 日付（必ず本日の日付を「YYYY年M月D日現在」形式で上書き）
    # 本文トップの「日現在」を含む段落を優先的に特定
    p_date = None
    for el in elems:
        if el.tag.endswith('}p'):
            txt = ''.join(t.text or '' for t in el.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))
            if '日現在' in txt or ('年' in txt and '現在' in txt):
                p_date = el
                break
    if p_date is None:
        p_date = find_para_with_text('年')
    if p_date is not None:
        set_para_elem_text(p_date, f"{today.year}年{today.month}月{today.day}日現在")

    # P[4]: 氏名
    p_name = find_para_with_text('氏名')
    if p_name is not None:
        set_para_elem_text(p_name, f"氏名\u3000{name}")

    # ■職務要約 の直後の空段落に要約を書く
    p_summary_head = find_para_with_text('■職務要約')
    if p_summary_head is not None:
        summary = d.get('workSummary', '')
        next_el = p_summary_head.getnext()
        if next_el is not None and next_el.tag.endswith('}p'):
            set_para_elem_text(next_el, summary)

    # ■活かせる経験 スキル箇条書き
    p_skills_head = find_para_with_text('■活かせる経験')
    if p_skills_head is not None:
        skills_text = d.get('skills', '')
        bullets = [s.strip() for s in skills_text.split('\n') if s.strip()]
        # 既存の '・' 段落を消して新しいものを挿入
        ref = p_skills_head
        next_el = ref.getnext()
        # 既存のスキル行を削除
        to_remove = []
        cur = next_el
        while cur is not None and cur.tag.endswith('}p'):
            txt = ''.join(t.text or '' for t in cur.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))
            if txt.strip() in ('・', '') or txt.strip().startswith('・'):
                to_remove.append(cur)
                cur = cur.getnext()
            else:
                break
        style_src = to_remove[0] if to_remove else p_skills_head
        for el in to_remove:
            el.getparent().remove(el)
        # 挿入（逆順にならないよう anchor を移動）
        last_inserted = p_skills_head
        for b in bullets:
            # AIが既に '・' で始めていても重複させない
            clean = b.lstrip('・•●◆▪️-–—\u3000 ').strip()
            if not clean:
                continue
            last_inserted = insert_paragraph_after(last_inserted, f"・{clean}", copy_style_from=style_src)
        if not bullets:
            insert_paragraph_after(p_skills_head, '・', copy_style_from=style_src)

    # ■職歴 — job sections を再構築
    # テンプレートには job section が 2 つある。3 つ以上は cloning
    p_shokureki = find_para_with_text('■職歴')
    if p_shokureki is not None:
        # ■職歴 以降の要素（■自己PR まで）を削除してから新規挿入
        p_jikopr = find_para_with_text('■自己PR')

        # ■職歴 と ■自己PR の間の要素を収集 → 削除
        collecting = False
        to_remove = []
        for el in list(body):
            if el is p_shokureki:
                collecting = True
                continue
            if el is p_jikopr:
                collecting = False
                break
            if collecting:
                to_remove.append(el)

        # 既存テンプレのジョブセクション (段落 + テーブル) をひとつ複製元として保持
        # 削除前に tbl を取得しておく
        tbl_template = None
        for el in to_remove:
            if el.tag.endswith('}tbl'):
                tbl_template = deepcopy(el)
                break
        p_company_template = None
        for el in to_remove:
            if el.tag.endswith('}p'):
                txt = ''.join(t.text or '' for t in el.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))
                if '株式会社' in txt or '年月' in txt:
                    p_company_template = deepcopy(el)
                    break

        for el in to_remove:
            el.getparent().remove(el)

        # ■自己PR の直前に jobs を挿入
        def insert_before(ref, el):
            ref.addprevious(el)

        # 各ジョブ挿入（逆順にしないと順序が崩れる → 順番通りに挿入、ただし基準を更新）
        if p_jikopr is None:
            p_jikopr = list(body)[-2]  # fallback

        prev_anchor = p_shokureki

        def make_para(text, style_src=None):
            if style_src is not None:
                el = deepcopy(style_src)
            else:
                el = OxmlElement('w:p')
            set_para_elem_text(el, text)
            return el

        def make_table_for_job(w, tbl_tmpl):
            """ジョブデータでテーブルを生成"""
            from docx.oxml.ns import qn
            tbl = deepcopy(tbl_tmpl)
            rows = tbl.findall(f'.//{{{qn("w:tr")[1:-1]}}}tr'.replace('{','').replace('}','')) if False else []
            # Use lxml directly
            ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            tr_list = tbl.findall(f'{{{ns}}}tr')
            if len(tr_list) < 3:
                return tbl

            def get_cells(tr):
                return tr.findall(f'{{{ns}}}tc')

            def set_cell_text(tc, text):
                """セル内の全段落をクリアし、複数行テキストを各段落として書く。
                元の段落数より行数が多ければ先頭段落を複製して追加、
                少なければ余分な段落を削除。最初の段落の書式を維持する。"""
                p_list = tc.findall(f'{{{ns}}}p')
                if not p_list: return
                lines = str(text if text is not None else '').split('\n')

                def write_para(p_el, line):
                    runs = p_el.findall(f'.//{{{ns}}}r')
                    for r in runs:
                        for t in r.findall(f'{{{ns}}}t'):
                            t.text = ''
                        # preserveSpace を有効にしておく
                    if runs:
                        t_el = runs[0].find(f'{{{ns}}}t')
                        if t_el is None:
                            t_el = OxmlElement('w:t')
                            runs[0].append(t_el)
                        t_el.text = line
                        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    else:
                        r_el = OxmlElement('w:r')
                        t_el = OxmlElement('w:t')
                        t_el.text = line
                        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                        r_el.append(t_el)
                        p_el.append(r_el)

                # 既存段落の数と行数を合わせる
                while len(p_list) < len(lines):
                    new_p = deepcopy(p_list[0])
                    # 新段落の中身はクリア
                    for r in new_p.findall(f'.//{{{ns}}}r'):
                        for t in r.findall(f'{{{ns}}}t'):
                            t.text = ''
                    p_list[-1].addnext(new_p)
                    p_list.append(new_p)
                # 行数が少なければ、余分な段落を空文字で埋める（削除すると書式が崩れるため残す）
                for i, p_el in enumerate(p_list):
                    line = lines[i] if i < len(lines) else ''
                    write_para(p_el, line)

            sy, sm = w.get('startYear', ''), w.get('startMonth', '')
            ey, em = w.get('endYear', ''), w.get('endMonth', '')
            current = w.get('isCurrent', False)
            etype  = w.get('employmentType', '正社員')
            job_content = w.get('jobContent', '')

            end_str = f"{ey}年{em}月" if not current and ey else ('現在')
            period_str = f"\n{sy}年{sm}月\n｜\n{end_str}"

            # row1: period + 雇用形態
            cells1 = get_cells(tr_list[1])
            if cells1:
                set_cell_text(cells1[0], period_str)
                if len(cells1) > 1:
                    set_cell_text(cells1[1], f"雇用形態：{etype}")
            # row2: period + 業務内容
            cells2 = get_cells(tr_list[2])
            if cells2:
                set_cell_text(cells2[0], period_str)
                if len(cells2) > 1:
                    set_cell_text(cells2[1], f"《業務内容》\n{job_content}")
            return tbl

        # 空行用テンプレ
        empty_p = OxmlElement('w:p')

        if not works:
            # ジョブなし → 空行
            ep = deepcopy(empty_p)
            prev_anchor.addnext(ep)
        else:
            # 各ジョブを [会社ヘッダー][空行][table][・事業内容][空行] の順で正しく挿入
            insert_anchor = p_shokureki
            for idx, w in enumerate(works):
                company = w.get('company', '')
                sy, sm = w.get('startYear', ''), w.get('startMonth', '')
                ey, em = w.get('endYear', ''), w.get('endMonth', '')
                current = w.get('isCurrent', False)
                end_str = f"{ey}年{em}月" if not current and ey else '現在'
                business = w.get('businessContent', '')

                # ① 会社名ヘッダー（番号付けはテンプレの自動リスト機能に任せ、本文のみを書く）
                header_text = f"{company}（{sy}年{sm}月〜{end_str}）"
                p_hdr = make_para(header_text, p_company_template)
                insert_anchor.addnext(p_hdr); insert_anchor = insert_anchor.getnext()

                # ② 空行
                insert_anchor.addnext(deepcopy(empty_p)); insert_anchor = insert_anchor.getnext()

                # ③ テーブル（期間・雇用形態・業務内容）
                if tbl_template is not None:
                    tbl = make_table_for_job(w, tbl_template)
                    insert_anchor.addnext(tbl); insert_anchor = insert_anchor.getnext()

                # ④ ・事業内容
                p_biz = make_para(f"・事業内容：{business}")
                insert_anchor.addnext(p_biz); insert_anchor = insert_anchor.getnext()

                # ⑤ 区切り空行（次ジョブとの間）
                insert_anchor.addnext(deepcopy(empty_p)); insert_anchor = insert_anchor.getnext()

    # ■自己PR
    p_jikopr2 = find_para_with_text('■自己PR')
    if p_jikopr2 is not None:
        selfpr = d.get('selfPr', '')
        next_el = p_jikopr2.getnext()
        if next_el is not None and next_el.tag.endswith('}p'):
            next2 = next_el.getnext()
            if next2 is not None and next2.tag.endswith('}p'):
                set_para_elem_text(next2, selfpr)
            else:
                set_para_elem_text(next_el, selfpr)

    # 末尾の「以上」段落を右寄せに
    for p in doc.paragraphs:
        if p.text.strip() == '以上':
            set_paragraph_alignment(p, 'right')

    # 全テキストを ＭＳ 明朝 に統一
    force_font(doc, 'ＭＳ 明朝')

    doc.save(dst)
    return {'label': '汎用 職務経歴書', 'filename': os.path.basename(dst), 'path': dst}


# ─────────────────────────────────────────────
# 3. ミラエール簡易履歴書.xlsx
# ─────────────────────────────────────────────

def gen_miraeru(d, out_dir, sid):
    from openpyxl import load_workbook

    src = os.path.join(TEMPLATES_DIR, 'ミラエール履歴書.xlsx')
    dst = os.path.join(out_dir, f"{safe_name(d.get('candidateName','候補者'))}_ミラエール履歴書_{stamp()}.xlsx")
    shutil.copy2(src, dst)
    wb = load_workbook(dst)
    ws = wb.active

    today = datetime.now()

    # 日付 (F3/J3/L3)
    ws['F3'] = today.year
    ws['J3'] = today.month
    ws['L3'] = today.day

    # ふりがな C4:N4
    ws['C4'] = d.get('candidateNameKana', '')
    # 氏名 C5:N6
    ws['C5'] = d.get('candidateName', '')
    # 生年月日 C7, F7, H7, L7, O7
    ws['C7'] = d.get('birthYear', '')
    ws['F7'] = d.get('birthMonth', '')
    ws['H7'] = d.get('birthDay', '')
    ws['L7'] = d.get('age', '')
    # N7:N8 は '男・女' ラベルのマージセル (top-left = N7) → 性別を上書き
    gender = d.get('gender', '')
    ws['N7'] = gender if gender else '男・女'
    # 住所ふりがな C9
    ws['C9'] = d.get('addressKana', '')
    # 携帯電話 — O9:Q9 はラベル「(携帯電話)：」が入ったマージセル。O9 (top-left) に値を追記する
    ws['O9'] = f" (携帯電話)：{d.get('mobile', d.get('phone', ''))}"
    # 住所 D10:Q10 の top-left は D10
    ws['D10'] = f"〒{d.get('zipCode','')}　{d.get('address','')}"
    # email C12
    ws['C12'] = d.get('email', '')

    # 学歴職歴 (rows 14-27, col B=年, C=月, D=内容)
    edu = d.get('educationHistory', [])
    works = d.get('workHistory', [])

    entries = []
    for e in edu:
        entries.append((e.get('year',''), e.get('month',''), e.get('content','')))

    for w in works:
        sy, sm = w.get('startYear',''), w.get('startMonth','')
        company = w.get('company','')
        etype   = w.get('employmentType','')
        entries.append((sy, sm, f"{company}　入社（{etype}）"))
        ey, em = w.get('endYear',''), w.get('endMonth','')
        if w.get('isCurrent'):
            entries.append(('', '', '現在に至る'))
        elif ey:
            entries.append((ey, em, f"{company}　退社"))

    for i, (y, m, c) in enumerate(entries[:14]):
        r = 14 + i
        ws.cell(r, 2, value=y)   # col B
        ws.cell(r, 3, value=m)   # col C
        ws.cell(r, 4, value=c)   # col D

    # 免許・資格 (rows 30-33, col D)
    licenses = d.get('licenses', [])
    for i, lic in enumerate(licenses[:4]):
        ws.cell(30 + i, 4, value=lic)

    # ミラエール専用: 希望勤務地 (B35 の後ろに追記)
    area = d.get('desiredArea', '')
    if area:
        cur = ws['B35'].value or ''
        # すでにある選択肢テキストを保持しつつ先頭に希望地を追記
        ws['B35'] = f"希望勤務地(1つ)：{area}"

    # 転居予定 (B36)
    reloc = d.get('relocationPlan', False)
    reloc_time = d.get('relocationTime', '')
    reloc_area = d.get('relocationArea', '')
    if reloc:
        ws['B36'] = f"転居予定(ある場合のみ記入)：時期（{reloc_time}）月頃、エリア・駅名（{reloc_area}）"

    # 署名欄 お名前 (E39)
    ws['E39'] = f"お名前：{d.get('candidateName','')}"

    wb.save(dst)
    return {'label': 'ミラエール 簡易履歴書', 'filename': os.path.basename(dst), 'path': dst}


# ─────────────────────────────────────────────
# 4. キャリアウィンク ES.xlsx
# ─────────────────────────────────────────────

def gen_carrierwink(d, out_dir, sid):
    from openpyxl import load_workbook

    src = os.path.join(TEMPLATES_DIR, 'キャリアウィンクES.xlsx')
    dst = os.path.join(out_dir, f"{safe_name(d.get('candidateName','候補者'))}_キャリアウィンクES_{stamp()}.xlsx")
    shutil.copy2(src, dst)
    wb = load_workbook(dst)
    ws = wb['エントリーシート']

    # 希望エリア D3
    ws['D3'] = d.get('desiredArea', '')
    # ふりがな D4 / E4
    kana = d.get('candidateNameKana', '')
    parts = kana.split() if kana else ['', '']
    ws['D4'] = parts[0] if parts else ''
    ws['E4'] = parts[1] if len(parts) > 1 else ''
    # 氏名 D5 / E5
    name = d.get('candidateName', '')
    nparts = name.split() if name else ['', '']
    ws['D5'] = nparts[0] if nparts else ''
    ws['E5'] = nparts[1] if len(nparts) > 1 else ''
    # 生年月日 D6 / E6 / F6
    ws['D6'] = d.get('birthYear', '')
    ws['E6'] = d.get('birthMonth', '')
    ws['F6'] = d.get('birthDay', '')
    # 性別 D8
    ws['D8'] = d.get('gender', '')
    # メール D9
    ws['D9'] = d.get('email', '')
    # 郵便番号 D10 (前3桁) / E10 (後4桁)
    zip_code = (d.get('zipCode', '') or '').strip()
    if '-' in zip_code:
        zparts = zip_code.split('-', 1)
    else:
        digits = ''.join(ch for ch in zip_code if ch.isdigit())
        zparts = [digits[:3], digits[3:]] if digits else ['', '']
    ws['D10'] = zparts[0] if zparts else ''
    ws['E10'] = zparts[1] if len(zparts) > 1 else ''
    # 住所 D11
    ws['D11'] = d.get('address', '')
    # 電話 D12
    ws['D12'] = d.get('mobile', d.get('phone', ''))
    # 学歴 D13
    ws['D13'] = d.get('lastEducation', '')
    # 面接希望日時 (rows 15-20: D=日付, E=時間帯)
    interview_dates = d.get('interviewDates', [])
    for i, iv in enumerate(interview_dates[:6]):
        r = 15 + i
        ws.cell(r, 4, value=format_interview_date_jp(iv.get('date', '')))   # D '4月28日'
        ws.cell(r, 5, value=format_interview_time(iv))  # E '終日可' or '10:00〜12:00'
    # 入社可能時期 D21
    ws['D21'] = d.get('availableFrom', '')
    # 備考 D23
    ws['D23'] = d.get('notes', '')

    wb.save(dst)
    return {'label': 'キャリアウィンク ES', 'filename': os.path.basename(dst), 'path': dst}


# ─────────────────────────────────────────────
# 5. エイジェック ES.xlsx
# ─────────────────────────────────────────────

def gen_agec(d, out_dir, sid):
    from openpyxl import load_workbook

    src = os.path.join(TEMPLATES_DIR, 'エイジェックES.xlsx')
    dst = os.path.join(out_dir, f"{safe_name(d.get('candidateName','候補者'))}_エイジェックES_{stamp()}.xlsx")
    shutil.copy2(src, dst)
    wb = load_workbook(dst)
    ws = wb['エントリーシート']

    # 氏名 D7
    ws['D7'] = d.get('candidateName', '')
    # フリガナ D8
    ws['D8'] = d.get('candidateNameKana', '')
    # メール D9
    ws['D9'] = d.get('email', '')
    # 性別 D10
    ws['D10'] = d.get('gender', '')
    # 電話 D11
    ws['D11'] = d.get('mobile', d.get('phone', ''))
    # 生年月日 D12
    by = d.get('birthYear',''); bm = d.get('birthMonth',''); bd = d.get('birthDay','')
    ws['D12'] = f"{by}/{bm}/{bd}" if by else ''
    # 年齢 D13
    ws['D13'] = d.get('age', '')

    # 職歴 3件 (rows 17-22)
    works = d.get('workHistory', [])
    job_rows = [(17, 18), (19, 20), (21, 22)]  # (勤続年数row, 業務内容row)
    for i, (tenure_row, content_row) in enumerate(job_rows):
        if i >= len(works): break
        w = works[i]
        tenure = w.get('tenure', '')
        content = w.get('jobContent', '')
        ws.cell(tenure_row, 5, value=tenure)    # E列
        ws.cell(content_row, 5, value=content)  # E列

    # 運転免許 E23
    licenses = d.get('licenses', [])
    driving = next((l for l in licenses if '運転' in l or '免許' in l), '')
    ws['E23'] = driving if driving else '無'
    # 入社可能日 E24
    ws['E24'] = d.get('availableFrom', '')
    # 勤務希望地 E25
    ws['E25'] = d.get('desiredArea', '')
    # 備考 E26：面接希望日時 + notes を合体
    remarks_parts = []
    interview_dates = d.get('interviewDates', []) or []
    if interview_dates:
        lines = ['【面接希望日時】']
        for iv in interview_dates[:6]:
            date_jp  = format_interview_date_jp(iv.get('date', ''))
            time_str = format_interview_time(iv)
            if not date_jp and not time_str:
                continue
            line = f"・{date_jp}　{time_str}".strip()
            lines.append(line)
        if len(lines) > 1:
            remarks_parts.append('\n'.join(lines))
    notes_txt = (d.get('notes', '') or '').strip()
    if notes_txt:
        remarks_parts.append(notes_txt)
    ws['E26'] = '\n\n'.join(remarks_parts) if remarks_parts else ''

    wb.save(dst)
    return {'label': 'エイジェック ES', 'filename': os.path.basename(dst), 'path': dst}


# ─────────────────────────────────────────────
# エントリーポイント
# ─────────────────────────────────────────────

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(json.dumps({'error': 'input JSON path required'}))
        sys.exit(1)

    with open(sys.argv[1], encoding='utf-8') as f:
        data = json.load(f)

    session_id = data.get('sessionId', 'sess')
    out_dir    = data.get('outputDir', '/tmp/ca_output')
    os.makedirs(out_dir, exist_ok=True)

    recommendations = data.get('recommendations', [])
    generated = []

    try:
        generated.append(gen_resume(data, out_dir, session_id))
    except Exception as e:
        generated.append({'label': '汎用 履歴書', 'error': str(e)})

    try:
        generated.append(gen_career(data, out_dir, session_id))
    except Exception as e:
        generated.append({'label': '汎用 職務経歴書', 'error': str(e)})

    if 'ミラエール' in recommendations:
        try:
            generated.append(gen_miraeru(data, out_dir, session_id))
        except Exception as e:
            generated.append({'label': 'ミラエール 簡易履歴書', 'error': str(e)})

    if 'キャリアウィンク' in recommendations:
        try:
            generated.append(gen_carrierwink(data, out_dir, session_id))
        except Exception as e:
            generated.append({'label': 'キャリアウィンク ES', 'error': str(e)})

    if 'エイジェック' in recommendations:
        try:
            generated.append(gen_agec(data, out_dir, session_id))
        except Exception as e:
            generated.append({'label': 'エイジェック ES', 'error': str(e)})

    print(json.dumps(generated, ensure_ascii=False))
